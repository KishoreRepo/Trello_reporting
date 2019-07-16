from docx import Document
from docx.shared import Inches
import datetime

import io, ast, json

document = Document()

def readFile():
    with io.open("trello_json_download.json",'r',encoding='utf-8-sig') as data:
        return data.read()

readFileData = readFile()

readFileDataJson = json.loads(readFileData)

print(readFileDataJson["name"])

document.add_heading(readFileDataJson["name"], 0)

document.add_heading('Initiators: user1 ', 4)
document.add_heading('Documentation owner: user2', 4)
document.add_heading('PoC\'s: User3', 4)

itemsStore={}

for action in readFileDataJson["actions"]:
    try:
        card_name = action["data"]["card"]["name"]
        if action["data"]["listAfter"]!="":            
            if itemsStore.get(card_name,None)==None:
                itemsStore[card_name]=[datetime.datetime.strptime(action["date"], "%Y-%m-%dT%H:%M:%S.%fZ"),action["data"]["listAfter"]["name"]]
            else:
                if itemsStore[card_name][0]<datetime.datetime.strptime(action["date"], "%Y-%m-%dT%H:%M:%S.%fZ"):
                    itemsStore[card_name]=[datetime.datetime.strptime(action["date"], "%Y-%m-%dT%H:%M:%S.%fZ"),action["data"]["listAfter"]["name"]]
    except Exception as e:
        pass
    
groupCat={}

for item,value in itemsStore.items():
    if groupCat.get(value[1],None)==None:
        groupCat[value[1]]=[item]
    else:
        groupCat[value[1]].append(item)

print(groupCat)

for item,values in groupCat.items():
    document.add_paragraph(item, style='Intense Quote')
    for value in values:
        
        document.add_paragraph(
        value, style='List Number'
        )

document.save('demo.docx')   
